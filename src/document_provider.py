#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档抽象层
统一 TXT / DOCX 等格式的段落读写接口，新格式只需实现 DocumentProvider 即可接入
"""

import os
import re
from abc import ABC, abstractmethod
from typing import Optional


class DocumentProvider(ABC):
    """文档抽象基类，定义统一的段落读写接口"""

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.file_name = os.path.basename(file_path)
        self.base_name = os.path.splitext(self.file_name)[0]
        self.file_ext = os.path.splitext(self.file_name)[1].lower()

    @abstractmethod
    def read_paragraphs(self) -> list[str]:
        """读取所有段落文本，返回段落列表"""
        ...

    @abstractmethod
    def apply_and_save(self, modifications: dict[int, str], output_path: Optional[str] = None) -> int:
        """
        应用修改并保存文件

        Args:
            modifications: {段落序号: 修改后文本}
            output_path: 输出路径，None 则自动推断

        Returns:
            实际修改的段落数
        """
        ...

    @abstractmethod
    def infer_output_path(self) -> str:
        """推断输出文件路径（如 input_fixed.txt / input_fixed.docx）"""
        ...


class TxtDocumentProvider(DocumentProvider):
    """TXT 文档提供者"""

    def __init__(self, file_path: str):
        super().__init__(file_path)

    def read_paragraphs(self) -> list[str]:
        # 按 UTF-8 / GBK 依次尝试
        for encoding in ('utf-8', 'gbk'):
            try:
                with open(self.file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                # 按空行分段，同时保留单换行的段落
                return [p.strip() for p in content.split('\n')]
            except (UnicodeDecodeError, UnicodeError):
                continue
        raise RuntimeError(f"[错误] 无法解码文件：{self.file_path}")

    def apply_and_save(self, modifications: dict[int, str], output_path: Optional[str] = None) -> int:
        if not modifications:
            return 0

        paragraphs = self.read_paragraphs()
        modified_count = 0

        # 按段落序号替换（1-based）
        for idx, modified_text in modifications.items():
            if 1 <= idx <= len(paragraphs) and modified_text:
                paragraphs[idx - 1] = modified_text
                modified_count += 1

        save_path = output_path or self.infer_output_path()
        with open(save_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(paragraphs))

        return modified_count

    def infer_output_path(self) -> str:
        return os.path.join(
            os.path.dirname(self.file_path),
            self.base_name + '_fixed.txt',
        )


class DocxDocumentProvider(DocumentProvider):
    """DOCX 文档提供者"""

    def __init__(self, file_path: str):
        super().__init__(file_path)
        from docx import Document
        self._doc = Document(file_path)

    def read_paragraphs(self) -> list[str]:
        return [p.text.strip() for p in self._doc.paragraphs]

    def apply_and_save(self, modifications: dict[int, str], output_path: Optional[str] = None) -> int:
        if not modifications:
            return 0

        modified_count = 0
        for idx, para in enumerate(self._doc.paragraphs, 1):
            if idx in modifications and modifications[idx]:
                _replace_para_text_with_refs(para, modifications[idx])
                modified_count += 1

        save_path = output_path or self.infer_output_path()
        self._doc.save(save_path)

        return modified_count

    def infer_output_path(self) -> str:
        return os.path.join(
            os.path.dirname(self.file_path),
            self.base_name + '_fixed.docx',
        )


def _replace_para_text_with_refs(para, text: str):
    """替换段落文本，将参考文献标记 [X] 设置为上标格式"""
    for run in para.runs:
        run.text = ''

    pattern = r'(\[\d+\])'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        run = para.add_run()
        run.text = part
        if re.match(r'^\[\d+\]$', part):
            run.font.superscript = True


def create_provider(file_path: str) -> DocumentProvider:
    """根据文件扩展名自动创建对应的 DocumentProvider"""
    ext = os.path.splitext(file_path)[1].lower()
    providers = {
        '.txt': TxtDocumentProvider,
        '.docx': DocxDocumentProvider,
    }
    cls = providers.get(ext)
    if cls is None:
        raise ValueError(f"[错误] 不支持的文件格式：{ext}，支持的格式：{', '.join(providers.keys())}")
    return cls(file_path)


__all__ = ['DocumentProvider', 'TxtDocumentProvider', 'DocxDocumentProvider', 'create_provider']

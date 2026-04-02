#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX 文件编辑器
逐段处理 DOCX 文件，直接应用语法检查结果
"""

import os
import re
from typing import Tuple
from docx import Document
from ollama_processor import chat_streaming
from config import DEFAULT_MODEL, PROCESSOR_CONFIGS
from file_io import save_to_txt
from content_splitter import is_table_of_contents, is_chapter_title


def is_ref_section(text: str) -> bool:
    """
    判断是否是参考文献章节

    Args:
        text: 待判断的文本

    Returns:
        bool: 是否是参考文献章节
    """
    text = text.strip()
    references_patterns = [
        r'^参考文献',  # 参考文献
        r'^参考文献\s*\[\d+\]',  # 参考文献 [1]
        r'^References$',  # References
        r'^References\s*\[\d+\]',  # References [1]
        r'^BIBLIOGRAPHY',  # BIBLIOGRAPHY
    ]
    for pattern in references_patterns:
        if re.match(pattern, text, re.IGNORECASE):
            return True
    return False


def _skip_para(text: str) -> bool:
    """
    判断段落是否应该跳过

    Args:
        text: 段落文本

    Returns:
        bool: 是否跳过
    """
    text_stripped = text.strip()

    # 1. 跳过空段落
    if not text_stripped:
        return True

    # 2. 跳过过短的段落（少于50个字符）
    if len(text_stripped) < 50:
        return True

    # 3. 跳过纯 LaTeX 公式段落
    # 检查段落是否主要由 LaTeX 公式组成（以 $ 或 \(\) 或 \[\] 包围）
    latex_pattern = r'^\s*([$\[\]\\].*[$\[\]\\]|\\[a-zA-Z]+\{.*\}|\\[a-zA-Z]+\s)+\s*$'
    if re.match(latex_pattern, text_stripped):
        return True

    # 4. 跳过目录内容
    if is_table_of_contents(text):
        return True

    # 5. 跳过章节标题
    if is_chapter_title(text):
        return True

    return False


def check_paragraph(paragraph_text: str) -> Tuple[str, str]:
    """
    检查单个段落,返回 Ollama 的原始输出
    复读检测和安全阈值由 Ollama 客户端自动处理

    Args:
        paragraph_text: 段落文本

    Returns:
        Tuple[str, str]: (Ollama 原始输出, 段落文本)
    """
    config = PROCESSOR_CONFIGS['grammar']

    # 构建提示词
    prompt = config['prompt_template'].format(content=paragraph_text)

    # 构建消息
    messages = [config['role'], {'role': 'user', 'content': prompt}]

    # 调用 Ollama(流式显示),复读检测由客户端自动处理
    result = chat_streaming(
        model=DEFAULT_MODEL,
        messages=messages
    ).strip()

    return result, paragraph_text


def _parse_modified_text(ollama_result: str) -> str:
    """
    从 Ollama 输出中解析修改后的文本

    Ollama 返回格式(由 GRAMMAR_CHECKER_PROMPT 定义):
    #### 第一部分：修改后的文本
    ###MODIFIED_TEXT###
    {修改后的文本内容}

    #### 第二部分：修改说明
    ###MODIFIED_DESCRIPTION###
    {修改说明}

    Args:
        ollama_result: Ollama 返回的完整结果

    Returns:
        str: 修改后的文本，如果没有修改则返回空字符串
    """
    if not ollama_result:
        return ""

    # 检查是否有 ###MODIFIED_TEXT### 分隔符
    if "###MODIFIED_TEXT###" not in ollama_result:
        return ""

    # 提取 MODIFIED_TEXT 之后的第一段内容
    parts = ollama_result.split("###MODIFIED_TEXT###")
    if len(parts) <= 1:
        return ""

    modified_part = parts[1].strip()

    # 按换行符分割，取第一个非空段落
    lines = modified_part.split('\n')
    for line in lines:
        if line.strip():
            return line.strip()

    return ""

def _infer_output_path(input_docx: str) -> str:
    """
    根据输入 DOCX 文件路径推断输出 DOCX 文件路径
    
    Args:
        input_docx: 输入 DOCX 文件路径
        
    Returns:
        str: 输出 DOCX 文件路径
    """
    input_dir = os.path.dirname(input_docx)
    base_name = os.path.splitext(os.path.basename(input_docx))[0]
    return os.path.join(input_dir, base_name + '_fixed.docx')

def apply_modifications(input_docx: str, modifications) -> int:
    """
    逐段处理 DOCX 文件，应用语法检查修改
    """
    if not modifications:
        print(f"[错误] 未找到任何修改记录：")
        return 0
    
    print(f"[信息] 解析到 {len(modifications)} 个段落修改记录")
    
    # 加载 DOCX 文件
    doc = Document(input_docx)
    modified_count = 0
    
    # 应用修改 - 直接按段落序号替换
    for idx, para in enumerate(doc.paragraphs, 1):
        if idx in modifications:
            # 获取修改后的文本
            modified_text = modifications[idx]

            # 直接替换段落，参考文献标记会自动设置为上标
            _replace_para_text_with_refs(para, modified_text)
            modified_count += 1

    output_docx = _infer_output_path(input_docx)
    # 保存修改后的文档
    doc.save(output_docx)
    
    print(f"处理完成！从 TXT 文件应用了 {modified_count} 个段落修改")
    print(f"保存: {output_docx}")

    return modified_count
    
    

def process_docx(input_docx: str, output_txt: str | None = None) -> Tuple[int, int]:
    """
    逐段处理 DOCX 文件，应用语法检查修改

    Args:
        input_docx: 输入 DOCX 文件路径
        output_txt: 语法检查结果 TXT 文件路径（可选），保存完整的修改建议

    Returns:
        Tuple[int, int]: (处理的总段落数, 修改的段落数)
    """
    if not os.path.exists(input_docx):
        print(f"[错误] 输入文件不存在：{input_docx}")
        return 0, 0

    doc = Document(input_docx)
    total_paragraphs = len(doc.paragraphs)
    skipped_count = 0

    # 找到最后一个"参考文献"段落的位置
    last_references_idx = -1
    for idx, para in enumerate(doc.paragraphs):
        if is_ref_section(para.text.strip()):
            last_references_idx = idx

    # 确定处理的段落范围（参考文献段落之前）
    process_until = last_references_idx if last_references_idx != -1 else total_paragraphs

    # 初始化 txt 文件（如果指定）
    if output_txt:
        save_to_txt("", output_txt, title="", mode='w')


    modifications = {}

    # 逐段检查并立即应用修改
    for idx, para in enumerate(doc.paragraphs[:process_until], 1):
        para_text = para.text.strip()

        # 检查是否应该跳过该段落
        if _skip_para(para_text):
            continue

        print(f"\r[段落 {idx} / {total_paragraphs}] 开始检查...",end='',flush=True)
        ollama_result, original = check_paragraph(para_text)

        # 保存到 txt 文件(只添加分隔线格式)
        if output_txt and ollama_result:
            _save_mods_to_txt(output_txt, idx, original, ollama_result)

        # 直接从 Ollama 输出解析修改后的文本
        corrected_text = _parse_modified_text(ollama_result)
        modifications[idx] = corrected_text
        
    modified_count = apply_modifications(input_docx, modifications)
    output_docx = _infer_output_path(input_docx)
    # 保存修改后的文档
    doc.save(output_docx)
    # 完成 txt 文件
    if output_txt:
        summary = (
            f"处理完成\n"
            f"总段落数：{total_paragraphs}\n"
            f"处理段落：{modified_count}\n"
            f"跳过段落：{total_paragraphs - modified_count}\n"
        )
        save_to_txt(summary, output_txt, title="=" * 80, mode='a')

    # 处理完成，输出结果
    print(f"处理完成！共处理 {total_paragraphs} 个段落，修改 {modified_count} 个")
    print(f"保存: {output_docx}")
    if output_txt:
        print(f"结果: {output_txt}")

    return total_paragraphs, modified_count


def _save_mods_to_txt(txt_file: str, para_idx: int, para_text: str, ollama_result: str):
    """
    将修改建议保存到 txt 文件（只添加格式化分隔线）

    Args:
        txt_file: txt 文件路径
        para_idx: 段落索引
        para_text: 原始段落文本
        ollama_result: Ollama 的原始输出
    """
    # 格式化内容，只添加分隔线
    content = f"{'=' * 60}\n"
    content += f"段落 {para_idx}\n"
    content += f"{'=' * 60}\n"
    content += f"原始内容：\n{para_text}\n\n"
    content += f"Ollama 输出：\n{ollama_result}\n\n"
    
    # 写入文件
    with open(txt_file, 'a', encoding='utf-8') as f:
        f.write(content)


def _parse_txt(txt_file: str) -> dict:
    """
    从 TXT 文件中解析段落修改记录
    
    TXT 格式(由 _save_modifications_to_txt 写入):
    ============================================================
    段落 N
    ============================================================
    原始内容：
    {原始段落文本}
    
    Ollama 输出：
    {Ollama 完整返回内容}
    
    Ollama 返回格式(由 GRAMMAR_CHECKER_PROMPT 定义):
    #### 第一部分：修改后的文本
    ###MODIFIED_TEXT###
    {修改后的文本内容}
    
    #### 第二部分：修改说明
    ###MODIFIED_DESCRIPTION###
    {修改说明}

    Args:
        txt_file: TXT 文件路径

    Returns:
        dict: 段落索引到修改后文本的映射 {para_idx: modified_text}
    """
    if not os.path.exists(txt_file):
        return {}

    modifications = {}
    import re
    
    with open(txt_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # 使用正则表达式匹配每个段落块
    # 模式: 60个= + 换行 + "段落 N" + 换行 + 60个= + ... + "Ollama 输出:" + ... + (到下一个段落或文件结尾)
    para_pattern = re.compile(
        r'={60}\s*\n段落\s+(\d+)\s*\n={60}\s*\n'
        r'原始内容：.*?\n\n'
        r'Ollama\s*输出：(.*?)(?=\n={60}\s*\n段落\s+\d+\s*\n={60}|\Z)',
        re.DOTALL
    )
    
    matches = para_pattern.finditer(content)
    
    for match in matches:
        para_idx = int(match.group(1))
        ollama_output = match.group(2).strip()
        
        # 从 Ollama 输出中解析修改后的文本
        modified_text = _parse_modified_text(ollama_output)
        
        if modified_text:
            modifications[para_idx] = modified_text
    
    return modifications


def apply_txt_to_docx(input_docx: str, txt_file: str) -> Tuple[int, int]:
    """
    从 TXT 文件解析修改记录并应用到 DOCX 文件
    直接按段落序号替换,无需匹配

    Args:
        input_docx: 输入 DOCX 文件路径
        txt_file: TXT 文件路径（包含修改记录）

    Returns:
        Tuple[int, int]: (处理的总段落数, 修改的段落数)
    """
    if not os.path.exists(txt_file):
        print(f"[错误] TXT 文件不存在：{txt_file}")
        return 0, 0
    
    # 解析 TXT 文件中的修改记录
    modifications = _parse_txt(txt_file)
    
    modified_count = apply_modifications(input_docx, modifications)
    
    return len(modifications), modified_count


def _replace_para_text_with_refs(para, text: str):
    """
    替换段落文本，并将参考文献标记 [X] 设置为上标格式
    写入前删除文内的空格

    Args:
        para: 段落对象
        text: 新的文本内容，可能包含 [1], [2] 等参考文献标记
    """
    import re

    # 删除文内空格（但保留中英文之间的分隔空格）
    # 这里删除所有空格，因为中文写作不需要空格
    text = re.sub(r' ', '', text)

    # 清空所有 run
    for run in para.runs:
        run.text = ''

    # 使用正则表达式分割文本，分离参考文献标记
    # 匹配 [数字] 格式的参考文献标记
    pattern = r'(\[\d+\])'
    parts = re.split(pattern, text)

    # 遍历分割结果，创建多个 run
    for part in parts:
        if not part:  # 跳过空字符串
            continue

        run = para.add_run()
        run.text = part

        # 如果是参考文献标记，设置为上标
        if re.match(r'^\[\d+\]$', part):
            run.font.superscript = True


__all__ = [
    'check_paragraph',
    'process_docx',
    'apply_txt_to_docx',
    '_parse_modified_text',
    '_replace_para_text_with_refs',
]

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件读写模块
处理所有文件的读取和保存操作
"""

import os


def read_file_content(file_path: str) -> str:
    """
    读取文件内容（支持 TXT 和 DOCX）

    Args:
        file_path: 文件路径

    Returns:
        str: 文件内容
    """
    file_ext = os.path.splitext(file_path)[1].lower()

    if file_ext == '.txt':
        return read_txt_content(file_path)
    elif file_ext == '.docx':
        return read_docx_content(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {file_ext}，仅支持 .txt 和 .docx")


def read_txt_content(file_path: str) -> str:
    """读取 TXT 文件内容"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding='gbk') as f:
            return f.read()


def read_docx_content(file_path: str) -> str:
    """
    读取 DOCX 文件内容

    Args:
        file_path: DOCX 文件路径

    Returns:
        str: 文件内容（纯文本）
    """
    try:
        from docx import Document

        doc = Document(file_path)

        # 按顺序提取所有内容（段落和表格）
        content_lines = []

        for element in doc.element.body:
            if element.tag.endswith('p'):  # 段落
                for para in doc.paragraphs:
                    if para._element == element:
                        if para.text.strip():
                            content_lines.append(para.text)
                        break
            elif element.tag.endswith('tbl'):  # 表格
                for table in doc.tables:
                    if table._element == element:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                content_lines.append(' | '.join(row_text))
                        content_lines.append('')  # 表格后空一行
                        break

        content = '\n'.join(content_lines)

        # 保存为 TXT 文件（与输入文件同名）
        txt_path = os.path.splitext(file_path)[0] + '.txt'
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(content)

        return content
    except ImportError:
        print("[错误] 未安装 python-docx 库，请运行: pip install python-docx")
        raise
    except Exception as e:
        print(f"[错误] 读取 DOCX 文件失败: {e}")
        raise


def validate_file_format(file_path: str) -> bool:
    """
    验证文件格式是否支持

    Args:
        file_path: 文件路径

    Returns:
        bool: 是否支持
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    if file_ext not in ['.txt', '.docx']:
        print(f"[错误] 不支持的文件格式: {file_ext}，仅支持 .txt 和 .docx，当前文件：{file_path}")
        return False
    return True


def save_to_txt(content: str, output_file: str, title: str = "学术风格摘要", mode: str = 'w'):
    """
    保存内容到 txt 文件

    Args:
        content: 要保存的内容
        output_file: 输出文件路径
        title: 文件标题（默认：学术风格摘要）
        mode: 文件打开模式，'w'为覆盖写入，'a'为追加写入（默认：'w'）

    Returns:
        None
    """
    with open(output_file, mode, encoding='utf-8') as f:
        f.write(f"{title}\n")
        f.write("=" * 60 + "\n")
        f.write(content)
        f.write("=" * 60 + "\n")
        f.flush()  # 立即刷新到磁盘


__all__ = [
    'read_file_content',
    'validate_file_format',
    'save_to_txt',
]
